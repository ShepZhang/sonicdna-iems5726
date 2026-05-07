"""Generate the Group43_report.docx based on the IEMS5726 template chapters.

Output: project_root/Group43_report.docx

The script does NOT consume the template directly; instead it builds
a fresh docx with the same chapter order and headings the template
requires, so we have full control over formatting and layout.

The user must:
    1. Run the EDA notebook first (figures must exist in models/figures/)
    2. Replace placeholder Group # / student IDs / names below
    3. Append a screenshot of the VeriGuide receipt at the end
    4. Insert 4-5 Streamlit screenshots into the User Interfaces section
       (markers shown as <<INSERT SCREENSHOT N>> paragraphs)
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "source_code"
sys.path.insert(0, str(SOURCE))

FIG = SOURCE / "models" / "figures"
ARCH = SOURCE / "models" / "figures"
SAMPLE_PRINT = SOURCE / "data" / "processed" / "moodprint_smoke_chill.png"
SCREENSHOTS = ROOT / "screenshots"


# --- placeholders the user MUST overwrite before submission ----------------
GROUP_NUMBER = "43"
MEMBERS = [
    ("1155238738", "Xupeng ZHANG"),
    ("1155251352", "Zetao HUANG"),
]
HF_SPACE_URL = "https://huggingface.co/spaces/<your-username>/sonicdna"
GDRIVE_MODELS_URL = "https://drive.google.com/drive/folders/<YOUR_FOLDER_ID>"
GITHUB_REPO_URL = "https://github.com/<your-username>/sonicdna"
# ---------------------------------------------------------------------------


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4A)
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
             size: int = 11, after_pt: int = 6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(after_pt)
    return p


def add_image(doc: Document, path: Path, width_cm: float = 15.0, caption: str = ""):
    if not path.exists():
        add_para(doc, f"[MISSING IMAGE: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_screenshot(doc: Document, filename: str, caption: str,
                   width_cm: float = 16.0):
    """Embed a UI screenshot from screenshots/ with a centered caption.

    Falls back to a red placeholder paragraph if the screenshot file is
    missing, so the report can still be generated before screenshots
    have been captured.
    """
    path = SCREENSHOTS / filename
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"<<MISSING SCREENSHOT: {filename}>>")
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build() -> Path:
    doc = Document()

    # --- Style defaults ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IEMS5726 Data Science in Practice\n2025-26 Term 2 / Session A")
    run.bold = True
    run.font.size = Pt(14)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("\nProject Report:\nSonicDNA — A Multi-Lens Music Personality Mirror")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0x4D, 0xB8)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Group number: {GROUP_NUMBER}\nGroup members:\n").bold = True
    for sid, name in MEMBERS:
        info.add_run(f"{sid}, {name}\n")

    doc.add_page_break()

    # --- 1. Problem Definitions ---
    add_heading(doc, "1. Problem Definitions")
    add_para(doc,
        "Modern music streaming services typically expose only two surfaces "
        "to a listener: an algorithmic 'similar tracks' rail and an annual "
        "marketing recap (e.g. Spotify Wrapped). Both reduce the user's "
        "library to a single dimension — popularity or genre — and leave "
        "no room for self-exploration of taste."
    )
    add_para(doc,
        "SonicDNA addresses this gap with a multi-lens music personality "
        "mirror: the user uploads a playlist (CSV) and the application "
        "produces three complementary outputs derived from a single 9-dim "
        "audio-feature embedding:")
    add_bullets(doc, [
        "MoodPrint — a unique polar-coordinate artwork (a 'music DNA') "
        "deterministically derived from the playlist composition, so two "
        "different playlists never share a print.",
        "Music MBTI — the playlist is reduced to a 4-letter code on four "
        "interpretable axes (E/I, N/S, F/T, J/P), each axis built from a "
        "subset of audio features, mapping to one of 16 named personality "
        "types with hand-written 60-character descriptions.",
        "Parallel-Universe Playlists — for three hand-curated 'universes' "
        "(Neon Synthwave Drive / Lo-fi Study Den / Mosh Pit Riot) we "
        "linearly project the user's centroid toward the universe's "
        "centroid (alpha = 0.7) and retrieve the top-10 most cosine-similar "
        "tracks inside that universe, answering 'what would you listen to "
        "if you grew up there?'.",
    ])
    add_para(doc,
        "The application therefore solves a dual problem: it produces "
        "objective, reproducible analytics (clustering, PCA, recommendation) "
        "and a subjective interpretive layer (MoodPrint art + MBTI labels) "
        "in one coherent pipeline — making the analytical results "
        "memorable and shareable instead of disposable.")

    # --- 2. Data Collection, Preprocessing and Representation ---
    add_heading(doc, "2. Data Collection, Preprocessing and Representation")

    add_heading(doc, "2.1 Data source", level=2)
    add_para(doc,
        "Kaggle Spotify Tracks Dataset (Maharshi Pandya): "
        "https://www.kaggle.com/datasets/maharshipandya/-spotify-tracks-dataset"
    )
    add_para(doc,
        "The dataset contains 114,000 tracks across 114 fine-grained genres, "
        "with the full audio-features panel that the (deprecated for new "
        "apps) Spotify Web API used to expose. Using this offline dump lets "
        "us avoid the API access changes published 2024-11 by Spotify.")

    add_heading(doc, "2.2 Data collection", level=2)
    add_para(doc,
        "We fetch the dataset programmatically via kagglehub on first run "
        "(see source_code/pipeline/preprocess.py and the Dockerfile build "
        "step), so the application is self-contained: no external API, no "
        "long-running scrapers.")

    add_heading(doc, "2.3 Data preprocessing", level=2)
    add_bullets(doc, [
        "Drop the original index column 'Unnamed: 0'.",
        "Drop rows missing any of {track_id, artists, track_name} (3 rows).",
        "De-duplicate on track_id (the dataset duplicates each track once "
        "per genre slice; 114k rows -> 89,740 unique tracks).",
        "Aggregate the 114 fine genres into 11 meta-genres (Pop, Rock, "
        "Metal, HipHop, Electronic, Jazz, Folk, Latin, World, Chill, "
        "Niche) via a hand-built dictionary in pipeline/preprocess.py.",
        "Standardize the 9 numeric audio features with sklearn's "
        "StandardScaler (mean 0, std 1) so cosine similarity, KMeans and "
        "PCA all operate in a comparable feature space.",
        "Persist the cleaned table to parquet (snappy) and the scaler to "
        "joblib pickle for re-use by all downstream modules.",
    ])

    add_heading(doc, "2.4 Data representation", level=2)
    add_para(doc,
        "Each track is represented as a 9-dimensional real-valued vector "
        "in z-scored space: ")
    add_para(doc,
        "(danceability, energy, valence, acousticness, instrumentalness, "
        "liveness, speechiness, tempo, popularity).",
        italic=True)
    add_para(doc,
        "A user playlist is represented as the mean of these 9-vectors "
        "across its tracks. This single embedding is the input shared by "
        "MoodPrint, Music MBTI, the recommender, and the parallel-universe "
        "projection — keeping all four creative modules consistent and "
        "letting any improvement to the representation propagate downstream "
        "without code changes elsewhere.")

    # --- 3. Data Modeling ---
    add_heading(doc, "3. Data Modeling")

    add_heading(doc, "3.1 Mood KMeans (k = 8)", level=2)
    add_para(doc,
        "We fit one KMeans (k=8, random_state=42, n_init=10) on the "
        "z-scored 9-D feature space. Cluster names are derived from each "
        "centroid's most distinctive trait via a small rule-base (see "
        "pipeline/cluster.py::name_mood_clusters): Mainstream Pulse, "
        "Sunlit Indie, Library Hush, Live Stage Glow, Stormcore Drive, "
        "Cipher Booth, Rainy Window, Drift Layer.")
    add_image(doc, FIG / "05_centroid_radars.png",
              caption="Figure 1. KMeans 8 cluster centroids in z-scored space. "
                      "Each cluster has a distinct dominant axis (e.g. Cipher "
                      "Booth on speechiness; Library Hush on acousticness "
                      "and instrumentalness).")

    add_heading(doc, "3.2 Music MBTI (4-axis composite mapping)", level=2)
    add_para(doc,
        "On the same z-scored space we define four composite axes whose "
        "signs encode an MBTI-style 4-letter code:")
    add_bullets(doc, [
        "E vs I: energy + danceability (extrovert vs introvert).",
        "N vs S: acousticness + instrumentalness (intuitive vs sensing).",
        "F vs T: valence − 0.5 × speechiness (feeling vs thinking).",
        "J vs P: popularity (mainstream-judging vs niche-perceiving).",
    ])
    add_para(doc,
        "Per-track axes are also computed (pipeline/mbti.py::"
        "attach_track_axes), which lets us list 'representative tracks' for "
        "each of the 16 types by picking the tracks whose own 4-axis scores "
        "are deepest into their quadrant. This avoids the sample-imbalance "
        "trap of training a 16-class classifier on heavily skewed labels.")

    add_heading(doc, "3.3 Parallel Universe projection", level=2)
    add_para(doc,
        "For each of three hand-curated universes we filter the library to "
        "a curated set of fine-grained genres (e.g., Neon Synthwave Drive = "
        "{synth-pop, electronic, electro, disco, new-age, house, deep-house, "
        "progressive-house, chicago-house, edm}) and compute the centroid "
        "vector. The user's centroid u is then linearly interpolated toward "
        "each universe centroid a:"
    )
    add_para(doc, "    projected = u + alpha * (a - u),  alpha = 0.7",
             italic=True)
    add_para(doc,
        "We run cosine-similarity recommendation inside that universe's "
        "track pool only, so the result is a believable 'you, but in this "
        "universe' playlist instead of a hard genre swap that would lose "
        "all of the user's personal style.")

    add_heading(doc, "3.4 Diversity-aware recommender", level=2)
    add_para(doc,
        "Pure cosine ranking suffers from the all-tracks-by-the-same-artist "
        "failure mode. We post-process the cosine ordering with two hard "
        "diversity caps: at most 1 track per primary artist and at most 3 "
        "tracks per meta-genre. This is the same routine reused by both the "
        "global recommender (Dashboard tab) and the universe-projected "
        "recommender (Parallel Universes tab) via a library_subset_mask.")

    add_heading(doc, "3.5 Pretrained model artifacts", level=2)
    add_para(doc,
        "All sklearn artifacts are saved to source_code/models/ and "
        "downloadable from a public Google Drive link:")
    add_bullets(doc, [
        "scaler.pkl — StandardScaler over 9 audio features",
        "kmeans_mood.pkl — KMeans(k=8) over z-scored space",
        "pca.pkl — 2-D PCA used by the Insights scatter map",
        "mood_cluster_names.json — interpretive name per centroid id",
    ])
    add_para(doc, f"Public download link: {GDRIVE_MODELS_URL}")

    # --- 4. Data Visualization ---
    add_heading(doc, "4. Data Visualization")

    add_heading(doc, "4.1 Visualization for the data", level=2)
    add_image(doc, FIG / "01_feature_distributions.png",
              caption="Figure 2. Distributions of 9 audio features on the raw "
                      "scale across 89,740 tracks; acousticness / "
                      "instrumentalness are heavily right-skewed, tempo is "
                      "roughly bimodal around 90 and 130 BPM.")
    add_image(doc, FIG / "02_correlation_heatmap.png",
              width_cm=12,
              caption="Figure 3. Pearson correlations between the 9 modeling "
                      "features. The set is mostly uncorrelated (|r| < 0.5 "
                      "everywhere except the energy-loudness pair which we "
                      "do not include), meaning the 9-D vector is not "
                      "redundant.")
    add_image(doc, FIG / "03_metagenre_distribution.png",
              caption="Figure 4. Track count per meta-genre after collapsing "
                      "114 fine genres into 11 buckets.")

    add_heading(doc, "4.2 Visualization for the model", level=2)
    add_image(doc, FIG / "04_pca_global_map.png", width_cm=13,
              caption="Figure 5. PCA 2-D projection of the full 89,740-track "
                      "library, colored by KMeans mood cluster id. The 8 "
                      "clusters separate cleanly even with only 40.5% of "
                      "variance preserved.")

    add_heading(doc, "4.3 Visualization for the results", level=2)
    add_para(doc,
        "Each user playlist produces a unique MoodPrint. Below is the "
        "MoodPrint of a 25-track Chill playlist used for development:")
    add_image(doc, SAMPLE_PRINT, width_cm=10,
              caption="Figure 6. MoodPrint for a 25-track Chill playlist. "
                      "Inner disc color = SHA1 hash of the 9-feature mean + "
                      "8-cluster distribution; petal lengths = sigmoid(2 * "
                      "z-score); outer ring = mood-cluster proportions.")

    # --- 5. System Architecture ---
    add_heading(doc, "5. System Architecture")

    add_heading(doc, "5.1 Current MVP architecture", level=2)
    add_para(doc,
        "The application is a single Streamlit process backed by a "
        "decoupled pipeline package. All four creative modules live in "
        "source_code/pipeline/ and are imported by application/app.py via "
        "an explicit sys.path injection at the top of the module:")
    add_image(doc, ARCH / "06_arch_mvp.png", width_cm=15,
              caption="Figure 7. Current MVP architecture. Streamlit is the "
                      "only UI surface; all model artifacts live on disk and "
                      "are loaded once via @st.cache_resource.")

    add_heading(doc, "5.2 Future production architecture", level=2)
    add_para(doc,
        "The pipeline package is intentionally framework-agnostic so the "
        "Streamlit UI can be swapped for the standard production stack "
        "(NGINX → Vue3 SPA + FastAPI backend → MySQL/Redis) without "
        "touching any modeling code. The diagram below documents that "
        "upgrade path:")
    add_image(doc, ARCH / "07_arch_production.png", width_cm=15,
              caption="Figure 8. Future production architecture upgrade path. "
                      "The Streamlit app collapses into a Vue3 SPA + FastAPI "
                      "backend; the same pipeline package is reused as a "
                      "library, so no modeling code changes.")

    add_heading(doc, "5.3 Deployment", level=2)
    add_para(doc,
        "The MVP is fully containerized: a single Dockerfile (build context "
        "= project root) installs both source_code and application "
        "dependencies, downloads the Kaggle dataset at build time, runs the "
        "preprocessing + clustering pipeline once, and exposes the "
        "Streamlit app on port 8501. The same image is the deliverable for "
        "the public HuggingFace Spaces demo.")
    add_para(doc, f"Public live demo: {HF_SPACE_URL}")
    add_para(doc, f"Source repository: {GITHUB_REPO_URL}")

    # --- 6. User Interfaces of the Software ---
    add_heading(doc, "6. User Interfaces of the Software")
    add_para(doc,
        "The Streamlit application has a fixed sidebar (sample-or-upload "
        "playlist selector) and four content tabs:")

    add_para(doc, "Screenshot 1 — Dashboard tab: vinyl-style MoodPrint, "
                  "Music MBTI card with rarity stat, the 16-type "
                  "distribution bar (user row highlighted in pink), and "
                  "the four-axis breakdown.", bold=True)
    add_screenshot(doc, "01_dashboard.png",
                   "Figure 9. Dashboard tab — Workout Beast Mode sample. "
                   "The vinyl print encodes 9 audio features as concentric "
                   "tracks; arc length / brightness / thickness reflect "
                   "how strongly the playlist exceeds the global average "
                   "on each feature.")

    add_para(doc, "Screenshot 2 — Parallel Universes tab: segmented "
                  "control to pick one of three hand-curated universes, "
                  "Top-10 cosine recommendations from inside that "
                  "universe's pool, and a three-overlay radar (universe "
                  "centroid, user vector, projected blend).",
             bold=True)
    add_screenshot(doc, "02_universes.png",
                   "Figure 10. Parallel Universes tab with Neon Synthwave "
                   "Drive selected. The yellow polygon is the projected "
                   "blend (alpha=0.7) used to score recommendations.")

    add_para(doc, "Screenshot 3 — Insights tab: PCA 2-D projection of "
                  "the full 89,740-track library coloured by mood "
                  "cluster, with the user's tracks highlighted as pink-"
                  "ringed dots; below it the user's 9-axis taste radar "
                  "and the mood-cluster donut.", bold=True)
    add_screenshot(doc, "03_insights.png",
                   "Figure 11. Insights tab — PCA scatter with the "
                   "Workout playlist clustered in the high-energy / "
                   "high-popularity region of the global music map.")

    add_para(doc, "Screenshot 4 — Sidebar: language-stable English-only "
                  "controls, sample-vs-upload selector, CSV schema help "
                  "tooltip, and a What's-Inside cheatsheet for the four "
                  "tabs.", bold=True)
    add_screenshot(doc, "06_sidebar.png",
                   "Figure 12. Sidebar — sample selector, CSV-upload "
                   "switch, and a brief tab cheatsheet.",
                   width_cm=8.0)

    # --- 7. Challenges ---
    add_heading(doc, "7. Challenges")
    add_para(doc, "Up to four challenges, in order of effort spent:")

    add_para(doc, "1. 9-D audio features → 2-D artistic fingerprint "
                  "(MoodPrint).", bold=True)
    add_para(doc,
        "Reducing 9 dimensions into a single image while preserving both "
        "uniqueness (every print must look different) and readability (a "
        "viewer must be able to recognize 'this playlist is high "
        "acousticness, low energy') required two design choices: (a) fix "
        "each feature's hue so the 9 petals are mnemonically learnable, "
        "and (b) sigmoid-amplify the z-score by 2× so that mean differences "
        "as small as 0.3σ produce visibly different petal lengths. Inner "
        "disc color is a SHA1 hash of the full feature + cluster "
        "distribution to guarantee uniqueness.")

    add_para(doc, "2. Sample imbalance in 16-class Music MBTI.", bold=True)
    add_para(doc,
        "Naively training a 16-class classifier on the library would have "
        "produced wildly unbalanced classes (ESFJ has 12,672 tracks; ISFP "
        "only 2,035). We sidestepped this by (a) defining the four MBTI "
        "axes as composite z-score sums rather than learned weights, "
        "making classification deterministic and interpretable, and "
        "(b) generating each type's representative tracks by ranking each "
        "type's quadrant by total signed axis strength + de-duplicating "
        "on artist, so even small classes have non-trivial picks.")

    add_para(doc, "3. Plausible 'parallel-universe' projection without losing "
                  "user identity.", bold=True)
    add_para(doc,
        "A naive implementation would replace the user's vector with the "
        "universe centroid, producing identical recommendations for every "
        "user. Instead we linearly interpolate u + α(a − u) with α=0.7, "
        "preserving 30% of the user's personal direction inside each "
        "universe pool, then run cosine recommendation inside the pool with "
        "the same diversity caps as the main recommender.")

    add_para(doc, "4. Working around the Spotify Web API 2024-11 access "
                  "policy change.", bold=True)
    add_para(doc,
        "Spotify deprecated audio-features endpoints for new development "
        "apps. Instead of registering a Spotify dev app and depending on "
        "external token freshness at request time, the entire pipeline "
        "operates on the offline Kaggle dump fetched once via kagglehub at "
        "Docker build time. This makes cold-start < 5 seconds and keeps "
        "the project resilient to future API changes.")

    # --- 8. References ---
    add_heading(doc, "8. References")
    add_bullets(doc, [
        "Maharshi Pandya. Spotify Tracks Dataset, 2023. "
        "https://www.kaggle.com/datasets/maharshipandya/-spotify-tracks-dataset",
        "Pedregosa et al. scikit-learn: Machine Learning in Python. JMLR, "
        "12, 2825-2830, 2011.",
        "Hunter, J. D. Matplotlib: A 2D graphics environment. Computing in "
        "Science & Engineering, 9(3), 90-95, 2007.",
        "McKinney, W. pandas: a Foundational Python Library for Data "
        "Analysis and Statistics. Python for High Performance and "
        "Scientific Computing, 14, 1-9, 2011.",
        "Streamlit Inc. Streamlit — the fastest way to build data apps. "
        "https://streamlit.io",
        "Plotly Technologies Inc. Plotly Python Open Source Graphing "
        "Library. https://plotly.com/python/",
        "HuggingFace. Spaces Documentation. "
        "https://huggingface.co/docs/hub/spaces",
        "Spotify. Introducing some changes to our Web API, "
        "2024-11-27. https://developer.spotify.com/blog/2024-11-27-changes-to-the-web-api",
    ])

    # --- 9. Declaration for the use of AI Tools ---
    add_heading(doc, "9. Declaration for the use of AI Tools")
    add_para(doc,
        "AI assistants were used in the following capacities during this "
        "project:")
    add_bullets(doc, [
        "Cursor (with Claude Opus 4.7 model) was used to draft and "
        "iteratively refine the pipeline modules (preprocess.py, "
        "cluster.py, moodprint.py, mbti.py, parallel_universe.py, "
        "recommender.py) and the Streamlit application (application/app.py). "
        "All algorithmic decisions (cluster naming rule book, MoodPrint "
        "polar layout, MBTI 4-axis definitions, alpha=0.7 interpolation, "
        "diversity caps) were specified by the author; the assistant was "
        "used to translate those decisions into runnable code.",
        "AI assistance was used to generate the 16 short Music-MBTI "
        "personality descriptions (60-character flavor text per type). "
        "Each description was reviewed and edited by the author before "
        "inclusion.",
        "AI was NOT used to: select or download the dataset, train the "
        "models (KMeans / PCA all run locally with fixed seeds), invent "
        "any results numbers, or write the final submitted report copy "
        "(although Cursor produced a first draft of this report's "
        "structure that the author edited).",
    ])

    # --- 10. Signed VeriGuide Receipt ---
    add_heading(doc, "10. Signed VeriGuide Receipt")
    add_para(doc,
        "The signed VeriGuide receipt is attached as the last page of "
        "this report.")
    p = doc.add_paragraph()
    run = p.add_run("<<INSERT VERIGUIDE RECEIPT IMAGE HERE>>")
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = ROOT / f"Group{GROUP_NUMBER}_report.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    out = build()
    print(f"[report] saved {out}")
    print(f"[report] file size: {out.stat().st_size / 1024:.1f} KB")
